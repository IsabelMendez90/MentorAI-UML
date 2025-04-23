import streamlit as st
from openai import OpenAI
import json
import pandas as pd
import re
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Leer la API Key desde Streamlit Secrets
API_KEY = st.secrets["OPENROUTER_API_KEY"]
API_BASE = "https://openrouter.ai/api/v1"
MODEL_NAME = "deepseek/deepseek-r1:free"

# Instrucciones del sistema
INSTRUCCIONES_SISTEMA = """
Eres un asistente experto en diseño de productos como integrador de sistemas mecatrónicos. Tu tarea es ayudar a generar soluciones para funciones de productos. 
Cuando el usuario proporcione una función, y su contexto de aplicación, debes responder con 5 soluciones técnicas viables, diversas y contextualizadas. Usa un lenguaje técnico, claro y conciso.
"""

# Inicializar estados de sesión
if "contexto_general" not in st.session_state:
    st.session_state.contexto_general = {}

if "historial_funciones" not in st.session_state:
    st.session_state.historial_funciones = []

if "respuestas_soluciones" not in st.session_state:
    st.session_state.respuestas_soluciones = []

if "limpiar_input" not in st.session_state:
    st.session_state.limpiar_input = False

if "matriz_generada" not in st.session_state:
    st.session_state.matriz_generada = False

if "combinaciones_generadas" not in st.session_state:
    st.session_state.combinaciones_generadas = []

# Función para obtener respuesta del modelo

def obtener_respuesta_funcion(mensaje):
     try:
         mensajes = [
             {"role": "system", "content": INSTRUCCIONES_SISTEMA},
             {"role": "user", "content": mensaje}
         ]
         client = OpenAI(api_key=API_KEY, base_url=API_BASE)
         completion = client.chat.completions.create(
             model=MODEL_NAME,
             messages=mensajes,
             extra_headers={
                 "HTTP-Referer": "https://mentor-ai-morpho.streamlit.app/",  
                 "X-Title": "Mentor-AI Matriz Morfologica"
             }
         )
         if completion and completion.choices:
             return completion.choices[0].message.content
         else:
             return "⚠️ No se recibió una respuesta válida del modelo. Intenta nuevamente."
     except Exception as e:
         return f"❌ Error al conectarse con el modelo: {str(e)}"


# Configuración de página
st.set_page_config(page_title="Mentor-AI Matriz Morfológica", layout="wide")
st.title("🤖 Challenge Mentor AI - Matriz Morfológica")
st.markdown(
    "Creadores: Dra. J. Isabel Méndez Garduño & M.Sc. Miguel de J. Ramírez C., CMfgT ")
st.subheader("Asistente interactivo")
st.markdown("Este asistente te ayuda a generar una matriz morfológica con soluciones para cada función técnica de tu producto.")



# Paso 1: Formulario de contexto
with st.expander("📋 Completa el contexto general antes de continuar", expanded=True):
    col1, col2 = st.columns(2)
    contexto = col1.text_area("a) Contexto del socio formador")
    reto = col2.text_area("b) Reto específico")
    necesidades = st.text_area("c) Necesidades del usuario")
    tareas = st.text_area("d) Tareas del usuario")
    ficha = st.text_area("e) Ficha técnica del prototipo o servicio")

contexto_completo = all([contexto.strip(), reto.strip(), necesidades.strip(), tareas.strip(), ficha.strip()])

if contexto_completo:
    st.session_state.contexto_general = {
        "contexto": contexto,
        "reto": reto,
        "necesidades": necesidades,
        "tareas": tareas,
        "ficha": ficha,
    }

# Paso 2: Historial tipo chat
st.subheader("💬 Historial de funciones y soluciones generadas")
if st.session_state.historial_funciones:
    for i in range(len(st.session_state.historial_funciones)):
        st.chat_message("user").write(st.session_state.historial_funciones[i])
        st.chat_message("assistant").markdown(st.session_state.respuestas_soluciones[i])
else:
    st.info("Aquí aparecerán las funciones que ingreses y las soluciones que te sugiera Mentor AI.")

# Mostrar historial de combinaciones
if st.session_state.combinaciones_generadas:
    st.subheader("🧠 Combinaciones óptimas generadas:")
    for idx, combinacion in enumerate(st.session_state.combinaciones_generadas, 1):
        st.markdown(f"**Propuesta #{idx}:**")
        st.markdown(combinacion)


# Paso 3: Ingreso de funciones (solo si contexto completo)
st.divider()
st.subheader("➕ Ingresa una nueva función")

if not contexto_completo:
    st.warning("⚠️ Completa el contexto general antes de ingresar funciones.")
else:
    if st.session_state.limpiar_input:
        st.session_state.funcion_input = ""
        st.session_state.limpiar_input = False

    with st.form("form_funcion"):
        nueva_funcion = st.text_input("🔹 Escribe la función del producto", key="funcion_input")
        enviar = st.form_submit_button("Generar soluciones")
        if enviar and nueva_funcion.strip():
            contexto_info = st.session_state.contexto_general
            prompt_funcion = f"""
La función del producto es: {nueva_funcion.strip()}.

Aquí tienes el contexto general del proyecto para generar soluciones más apropiadas:

a) Contexto del socio formador:
{contexto_info['contexto']}

b) Reto específico:
{contexto_info['reto']}

c) Necesidades del usuario:
{contexto_info['necesidades']}

d) Tareas del usuario:
{contexto_info['tareas']}

e) Ficha técnica del prototipo:
{contexto_info['ficha']}

Con base en esta información, proporciona exactamente **5 soluciones técnicas posibles** para esta función.
Para cada una, sigue este formato:
- Escribe únicamente el **nombre/título** de la solución como encabezado.
- Luego agrega una breve descripción que contenga 4 secciones:
  - Tecnología:
  - Funcionamiento:
  - Ventajas:
  - Desventajas:

Al final, si es útil, agrega una sección de **Notas transversales**.
No uses numeración ni encabezados como \"Solución 1:\", solo el nombre de cada solución como título.
"""
            with st.spinner("🧠 Generando soluciones..."):
                respuesta = obtener_respuesta_funcion(prompt_funcion)
            st.session_state.historial_funciones.append(nueva_funcion.strip())
            st.session_state.respuestas_soluciones.append(respuesta)
            st.session_state.limpiar_input = True
            st.rerun()

# Matriz morfológica
st.divider()
if len(st.session_state.historial_funciones) >= 3:
    if st.button("📄 Generar Matriz Morfológica"):
        matriz_data = []
        secciones = ["tecnología", "funcionamiento", "ventajas", "desventajas"]

        for i, funcion in enumerate(st.session_state.historial_funciones):
            soluciones_texto = st.session_state.respuestas_soluciones[i]
            lineas = soluciones_texto.strip().split("\n")
            soluciones = []

            for idx, linea in enumerate(lineas):
                linea_clean = linea.strip()
                if not linea_clean:
                    continue
                if not any(sec in linea_clean.lower() for sec in secciones):
                    siguientes = lineas[idx+1:idx+5]  # buscar en las siguientes 4 líneas
                    if any(any(sec in s.lower() for sec in secciones) for s in siguientes):
                        titulo = re.sub(r'^[-•*\d+\s]*', '', linea_clean)
                        titulo = titulo.replace("**", "").strip()
                        soluciones.append(titulo)

            soluciones = (soluciones + ["—"] * 5)[:5]
            matriz_data.append([funcion] + soluciones)

        df_morfologica = pd.DataFrame(matriz_data, columns=[
            "Función", "Solución 1", "Solución 2", "Solución 3", "Solución 4", "Solución 5"
        ])
        st.session_state.df_morfologica = df_morfologica
        st.session_state.matriz_generada = True

if st.session_state.get("matriz_generada", False):
    st.subheader("🧩 Matriz Morfológica Generada")
    st.dataframe(st.session_state.df_morfologica, use_container_width=True)

    # Paso 5: Generar combinación por criterio
    st.divider()
    st.subheader("🎯 Propuestas de combinación según criterio")

    criterios_predefinidos = [
        "Tecnología avanzada",
        "Costo-beneficio",
        "Más barato",
        "Más realista",
        "Rápido de implementar",
        "Otro"
    ]

    col1, col2 = st.columns(2)
    criterio_default = col1.selectbox("Selecciona un criterio predefinido:", criterios_predefinidos)

    criterio_personalizado = ""
    if criterio_default == "Otro":
        criterio_personalizado = col2.text_input("Escribe tu propio criterio:")




    if st.button("🔍 Generar combinación óptima"):
        criterio_final = criterio_personalizado.strip() if criterio_default == "Otro" and criterio_personalizado.strip() else criterio_default

        funciones = st.session_state.historial_funciones
        opciones = []

        for i, soluciones_texto in enumerate(st.session_state.respuestas_soluciones):
            soluciones = []
            lineas = soluciones_texto.strip().split("\n")
            secciones = ["tecnología", "funcionamiento", "ventajas", "desventajas"]

            for idx, linea in enumerate(lineas):
                linea_clean = linea.strip()
                if not linea_clean:
                    continue
                if not any(sec in linea_clean.lower() for sec in secciones):
                    siguientes = lineas[idx+1:idx+5]
                    if any(any(sec in s.lower() for sec in secciones) for s in siguientes):
                        titulo = re.sub(r'^[-•*\d+\s]*', '', linea_clean)
                        titulo = titulo.replace("**", "").strip()
                        soluciones.append(titulo)

            soluciones = (soluciones + ["—"] * 5)[:5]
            opciones.append({"funcion": funciones[i], "opciones": soluciones})

        prompt = f"""
Tienes la siguiente matriz morfológica con funciones y 5 soluciones posibles por cada una. El criterio de evaluación es: "{criterio_final}".

Para cada función, selecciona una sola solución que sea la mejor bajo este criterio. Por cada selección, indica el número de solución elegida (por ejemplo S1, S2...), y justifica por qué la elegiste.

Después, resume la combinación en una sola línea así:
Criterio: {criterio_final} → [Función 1 (Sx)] + [Función 2 (Sx)] + ...

Estructura tu respuesta así:

Función: [nombre]
Solución seleccionada: [texto completo de la solución] (Sx)
Justificación: [una oración clara]

Resumen al final:
Criterio: {criterio_final} → [Función 1 [texto completo de la solución]] + [Función 2 [texto completo de la solución]] + ...
...

Funciones y opciones:
{json.dumps(opciones, indent=2, ensure_ascii=False)}
"""

        with st.spinner("🧠 Analizando combinación óptima..."):
            resultado_combinacion = obtener_respuesta_funcion(prompt)

        resumen_limpio = resultado_combinacion.replace("Resumen al final:\nCriterio seleccionado →", f"Criterio: {criterio_final} →")
        st.session_state.combinaciones_generadas.append(resumen_limpio)


def agregar_parrafo_con_formato(doc, texto):
    parrafo = doc.add_paragraph()
    cursor = 0
    while cursor < len(texto):
        if texto[cursor:cursor+2] == '**':
            cursor += 2
            fin = texto.find('**', cursor)
            if fin == -1:
                parrafo.add_run('**' + texto[cursor:])  # por si está mal cerrado
                break
            run = parrafo.add_run(texto[cursor:fin])
            run.bold = True
            cursor = fin + 2
        elif texto[cursor] == '*':
            cursor += 1
            fin = texto.find('*', cursor)
            if fin == -1:
                parrafo.add_run('*' + texto[cursor:])
                break
            run = parrafo.add_run(texto[cursor:fin])
            run.italic = True
            cursor = fin + 1
        else:
            fin = cursor
            while fin < len(texto) and texto[fin] != '*' and texto[fin:fin+2] != '**':
                fin += 1
            parrafo.add_run(texto[cursor:fin])
            cursor = fin
    return parrafo



# Mostrar historial de combinaciones
if st.session_state.combinaciones_generadas:
    st.subheader("🧠 Combinaciones óptimas generadas:")
    for idx, combinacion in enumerate(st.session_state.combinaciones_generadas, 1):
        st.markdown(f"**Propuesta #{idx}:**")
        st.markdown(combinacion)

    # 🔽 Mostrar botones de descarga SIEMPRE que haya combinaciones
    st.divider()
    st.subheader("📥 Descargar reportes")

    fecha_hora_actual = datetime.now().strftime("%Y%m%d-%H%M")

    doc = Document()
    doc.add_heading("Reporte Mentor-AI: Matriz Morfológica", level=1)
    doc.add_heading("📌 Contexto General", level=2)
    doc.add_paragraph(f"Contexto del socio formador: {st.session_state.contexto_general['contexto']}")
    doc.add_paragraph(f"Reto específico: {st.session_state.contexto_general['reto']}")
    doc.add_paragraph(f"Necesidades del usuario: {st.session_state.contexto_general['necesidades']}")
    doc.add_paragraph(f"Tareas del usuario: {st.session_state.contexto_general['tareas']}")
    doc.add_paragraph(f"Ficha técnica del prototipo: {st.session_state.contexto_general['ficha']}")

    doc.add_heading("🔧 Funciones y soluciones generadas", level=2)
    for i, (funcion, respuesta) in enumerate(zip(st.session_state.historial_funciones, st.session_state.respuestas_soluciones)):
        doc.add_heading(f"Función {i+1}: {funcion}", level=3)
        for linea in respuesta.split("\n"):
            if linea.strip():
                p = agregar_parrafo_con_formato(doc, linea.strip())
                p.style.font.size = Pt(11)

    doc.add_heading("🎯 Combinaciones óptimas generadas", level=2)
    for i, combinacion in enumerate(st.session_state.combinaciones_generadas):
        doc.add_paragraph(f"Propuesta #{i+1}:")
        for linea in combinacion.split("\n"):
            if linea.strip():
                p = agregar_parrafo_con_formato(doc, linea.strip())
                p.style.font.size = Pt(11)

    buffer_word = BytesIO()
    doc.save(buffer_word)
    buffer_word.seek(0)
    st.download_button(
        label="📄 Descargar Word",
        data=buffer_word,
        file_name=f"{fecha_hora_actual}-Reporte_M-Morfologica.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    if "df_morfologica" in st.session_state:
        buffer_excel = BytesIO()
        st.session_state.df_morfologica.to_excel(buffer_excel, index=False)
        buffer_excel.seek(0)
        st.download_button(
            label="📊 Descargar Excel",
            data=buffer_excel,
            file_name=f"{fecha_hora_actual}-Tabla_M-Morfologica.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
